"""Unified upload intake — classify files and extract text or route to OCR."""

from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass
from typing import Literal

from PIL import Image, UnidentifiedImageError

from services.document.pdf import DocumentExtractionError, extract_pdf_text

DEFAULT_MAX_FILE_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARS = 100_000

FileKind = Literal["image", "text"]


@dataclass(frozen=True)
class FileIntake:
    kind: FileKind
    filename: str
    content_type: str | None = None
    text: str | None = None
    image_bytes: bytes | None = None


class FileIntakeError(ValueError):
    def __init__(self, message: str = "The uploaded file could not be processed.") -> None:
        super().__init__(message)
        self.message = message


def _ext(filename: str) -> str:
    match = re.search(r"(\.[a-z0-9]{1,8})$", (filename or "").lower())
    return match.group(1) if match else ""


def _truncate(text: str) -> str:
    text = text.strip()
    if not text:
        return text
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text
    return f"{text[:MAX_EXTRACTED_CHARS]}\n[File content truncated for analysis]"


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise FileIntakeError("Could not decode this file as text.")


def _looks_binary(data: bytes) -> bool:
    if not data:
        return False
    sample = data[:4096]
    if b"\x00" in sample:
        return True
    printable = sum(32 <= b <= 126 or b in (9, 10, 13) for b in sample)
    return printable / max(len(sample), 1) < 0.75


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise FileIntakeError("Word document support is not available on this server.") from exc

    try:
        document = Document(io.BytesIO(data))
        parts = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
    except Exception as exc:
        raise FileIntakeError("The Word document could not be read.") from exc

    if not text:
        raise FileIntakeError("No readable text found in this Word document.")
    return _truncate(text)


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise FileIntakeError("Excel spreadsheet support is not available on this server.") from exc

    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
    except Exception as exc:
        raise FileIntakeError("The Excel spreadsheet could not be read.") from exc

    if not text:
        raise FileIntakeError("No readable text found in this spreadsheet.")
    return _truncate(text)


def _extract_csv(data: bytes) -> str:
    decoded = _decode_text(data)
    reader = csv.reader(io.StringIO(decoded))
    rows = [" | ".join(cell.strip() for cell in row if cell.strip()) for row in reader]
    rows = [row for row in rows if row]
    if not rows:
        raise FileIntakeError("No readable rows found in this CSV file.")
    return _truncate("\n".join(rows))


def _extract_json(data: bytes) -> str:
    decoded = _decode_text(data)
    try:
        payload = json.loads(decoded)
    except json.JSONDecodeError as exc:
        raise FileIntakeError("The JSON file is not valid.") from exc
    return _truncate(json.dumps(payload, indent=2, ensure_ascii=False))


def _try_image_intake(data: bytes, filename: str, content_type: str | None) -> FileIntake | None:
    try:
        with Image.open(io.BytesIO(data)) as img:
            img.verify()
        with Image.open(io.BytesIO(data)) as img:
            fmt = (img.format or "").upper()
    except (UnidentifiedImageError, OSError):
        return None

    if not fmt:
        return None

    mime = content_type or _IMAGE_FORMAT_TO_MIME.get(fmt, "application/octet-stream")
    return FileIntake(
        kind="image",
        filename=filename,
        content_type=mime,
        image_bytes=data,
    )


_IMAGE_FORMAT_TO_MIME = {
    "PNG": "image/png",
    "JPEG": "image/jpeg",
    "WEBP": "image/webp",
    "GIF": "image/gif",
    "BMP": "image/bmp",
    "TIFF": "image/tiff",
    "MPO": "image/jpeg",
}


def intake_file(
    data: bytes,
    *,
    filename: str = "upload",
    content_type: str | None = None,
    max_bytes: int = DEFAULT_MAX_FILE_BYTES,
) -> FileIntake:
    """Classify an upload as image (OCR path) or text (document path)."""
    if not data:
        raise FileIntakeError("The file is empty.")
    if len(data) > max_bytes:
        raise FileIntakeError("The file exceeds the maximum allowed size (10 MB).")

    name = filename or "upload"
    ext = _ext(name)
    mime = (content_type or "").split(";")[0].strip().lower()

    image = _try_image_intake(data, name, content_type)
    if image:
        return image

    text: str | None = None

    try:
        if ext == ".pdf" or mime == "application/pdf":
            text = extract_pdf_text(data, max_bytes=max_bytes)
        elif ext == ".docx" or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = _extract_docx(data)
        elif ext in {".xlsx", ".xlsm"} or mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            text = _extract_xlsx(data)
        elif ext == ".csv" or mime == "text/csv":
            text = _extract_csv(data)
        elif ext == ".json" or mime == "application/json":
            text = _extract_json(data)
        elif ext in {".txt", ".md", ".markdown", ".log", ".yaml", ".yml", ".xml", ".html", ".htm", ".rtf"} or mime.startswith("text/"):
            text = _truncate(_decode_text(data))
        elif ext == ".doc" or mime == "application/msword":
            raise FileIntakeError("Legacy .doc files are not supported. Save as .docx and try again.")
        elif ext == ".xls" or mime == "application/vnd.ms-excel":
            raise FileIntakeError("Legacy .xls files are not supported. Save as .xlsx and try again.")
        elif ext == ".ppt" or ext == ".pptx":
            raise FileIntakeError("PowerPoint files are not supported yet. Export to PDF or paste the text.")
        elif not _looks_binary(data):
            text = _truncate(_decode_text(data))
        else:
            raise FileIntakeError(
                "This file type could not be read. Try PDF, Word, Excel, CSV, plain text, or an image."
            )
    except DocumentExtractionError as exc:
        raise FileIntakeError(exc.message) from exc

    if not text:
        raise FileIntakeError("No readable content found in this file.")

    return FileIntake(
        kind="text",
        filename=name,
        content_type=content_type,
        text=text,
    )
